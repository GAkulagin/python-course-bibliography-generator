"""
Функции для генерации выходного файла с оформленным списком использованных источников.
"""
from __future__ import annotations
from abc import ABC, abstractmethod
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pylint: disable=E0611
from docx.shared import Pt, Mm


class BaseRenderer(ABC):
    """
        Базовый класс для создания word-файла
    """

    def __init__(self, rows: tuple[str, ...]):
        self.rows = rows

    @abstractmethod
    def render(self, path: Path | str) -> None:
        """
            Метод генерации Word-файла со списком использованных источников.

            :param Path | str path: Путь для сохранения выходного файла.
        """


class GOSTRenderer(BaseRenderer):

    def render(self, path: Path | str) -> None:

        document = Document()

        # стилизация заголовка
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        runner = paragraph.add_run("Список использованной литературы")
        runner.font.size = Pt(16)
        runner.bold = True

        # стилизация текста
        style_normal = document.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(13)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_normal.paragraph_format.keep_together = True

        for row in self.rows:
            # добавление источника
            document.add_paragraph(row, style="List Number")

        # сохранение файла Word
        document.save(path)


class APARenderer(BaseRenderer):

    def render(self, path: Path | str) -> None:
        document = Document()

        # стилизация заголовка
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = paragraph.add_run("References")
        runner.font.size = Pt(14)
        runner.bold = True

        # стилизация текста
        style_normal = document.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_normal.paragraph_format.space_after = Pt(12)
        style_normal.paragraph_format.first_line_indent = Mm(-10)
        style_normal.paragraph_format.keep_together = True

        for row in self.rows:
            # добавление источника
            document.add_paragraph(row, style="Normal")

        # сохранение файла Word
        document.save(path)
